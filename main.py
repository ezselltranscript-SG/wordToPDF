from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import uuid
import subprocess
import logging
import re
import tempfile
import shutil
from pathlib import Path
from docx import Document
from PyPDF2 import PdfWriter, PdfReader
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx.shared import Pt

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Crear directorios para archivos
UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")

# Asegurarse de que los directorios existan
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

app = FastAPI(
    title="Word to PDF Converter API",
    description="API sencilla para convertir documentos Word a PDF",
    version="1.0.0"
)

# Configurar CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/convert/", summary="Convertir documento Word a PDF")
async def convert_word_to_pdf(file: UploadFile = File(...), background_tasks: BackgroundTasks = None):
    """
    Convierte un documento Word (.docx) a formato PDF.
    
    - **file**: Archivo Word (.docx) a convertir
    
    Retorna el archivo PDF convertido.
    """
    # Verificar que el archivo sea un documento Word
    if not file.filename.endswith(('.docx', '.doc')):
        logger.warning(f"Archivo no válido: {file.filename}")
        raise HTTPException(status_code=400, detail="El archivo debe ser un documento Word (.docx o .doc)")
    
    # Generar nombres únicos para los archivos
    file_id = str(uuid.uuid4())
    input_filename = f"{file_id}_{file.filename}"
    input_path = UPLOAD_DIR / input_filename
    
    try:
        # Guardar el archivo subido
        with open(input_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        logger.info(f"Archivo guardado en {input_path}")
        
        # Modificar el documento para corregir los encabezados
        result = await modify_document_headers(str(input_path))
        
        if not result or not result[0]:
            logger.error(f"Error al modificar encabezados en {input_path}")
            raise HTTPException(status_code=500, detail="Error al procesar el documento")
        
        modified_docx, doc_base_code = result
        
        # Usar exactamente el nombre original del archivo sin modificaciones
        base_code = Path(file.filename).stem
        # Asegurarse de que no se modifique el formato (no convertir a minúsculas, etc.)
        logger.info(f"Código base del nombre del archivo: {base_code}")
        
        # Convertir a PDF usando LibreOffice
        pdf_filename = f"{Path(file.filename).stem}.pdf"
        output_pdf = await convert_to_pdf(modified_docx, str(OUTPUT_DIR))
        
        if not output_pdf:
            logger.error(f"Error al convertir {modified_docx}")
            raise HTTPException(status_code=500, detail="Error al convertir el documento")
        
        # Modificar el PDF para añadir encabezados correctos en cada página
        modified_pdf = await add_page_headers_to_pdf(output_pdf, base_code)
        
        if not modified_pdf:
            logger.error(f"Error al modificar encabezados en el PDF {output_pdf}")
            raise HTTPException(status_code=500, detail="Error al modificar encabezados en el PDF")
        
        logger.info(f"Conversión exitosa con encabezados modificados: {modified_pdf}")
        
        # Usar el PDF modificado como resultado final
        output_pdf = modified_pdf
        
        # Limpiar archivos temporales
        if background_tasks:
            def cleanup():
                for path in [input_path, modified_docx]:
                    if os.path.exists(str(path)):
                        try:
                            os.remove(str(path))
                            logger.info(f"Archivo temporal eliminado: {path}")
                        except Exception as e:
                            logger.error(f"Error al eliminar archivo temporal {path}: {str(e)}")
            
            background_tasks.add_task(cleanup)
        
        # Devolver el archivo PDF
        return FileResponse(
            path=output_pdf,
            media_type="application/pdf",
            filename=pdf_filename
        )
        
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        # Limpiar archivo temporal en caso de error
        if os.path.exists(input_path):
            os.remove(input_path)
        raise HTTPException(status_code=500, detail="Error al convertir el documento")

async def modify_document_headers(docx_path):
    """
    Modifica los encabezados del documento Word para que cada página tenga el formato correcto
    con Part1, Part2, Part3, etc.
    """
    try:
        # Extraer el nombre base del archivo
        original_filename = os.path.basename(docx_path).split('_', 1)[1] if '_' in os.path.basename(docx_path) else os.path.basename(docx_path)
        temp_dir = tempfile.mkdtemp()
        base_name = os.path.basename(docx_path)
        modified_docx = os.path.join(temp_dir, f"modified_{base_name}")
        
        # Abrir el documento original
        doc = Document(docx_path)
        
        # Extraer el código base del nombre del archivo
        base_code = None
        
        # Extraer el código base del nombre del archivo exactamente como aparece
        # Ejemplo: "062725-0620-b04-25.docx" -> "062725-0620-b04-25"
        base_code = os.path.splitext(base_name)[0]
        logger.info(f"Código base identificado: {base_code}")
        
        # Si no se encuentra un código base, usar un valor predeterminado
        if not base_code:
            base_code = "transcript"
            logger.warning(f"No se identificó código base, usando valor predeterminado: {base_code}")
        
        # ELIMINAR COMPLETAMENTE los encabezados de cada sección
        for section_idx, section in enumerate(doc.sections):
            part_number = section_idx + 1
            header = section.header
            
            # Eliminar todo el contenido del encabezado
            for paragraph in list(header.paragraphs):
                p = paragraph._element
                p.getparent().remove(p)
                paragraph._p = None
                paragraph._element = None
            
            # Añadir un párrafo vacío para mantener la estructura
            header.add_paragraph()
            
            logger.info(f"Eliminado encabezado para sección {part_number}")
        
        # Forzar Times New Roman 10 en todos los estilos
        try:
            from docx.enum.style import WD_STYLE_TYPE
            for style in doc.styles:
                if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
                    if style.font is not None:
                        style.font.name = 'Times New Roman'
                        style.font.size = Pt(10)
        except Exception as e:
            logger.warning(f"No se pudo modificar estilos globales: {e}")

        # Cambiar la fuente y tamaño manualmente en cada ejecución de texto
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        # Cambiar también en las tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)
        
        # Guardar el documento modificado
        doc.save(modified_docx)
        logger.info(f"Documento con encabezados eliminados guardado en: {modified_docx}")
        
        # Devolver el documento modificado y el código base
        return modified_docx, base_code
        
    except Exception as e:
        logger.error(f"Error al modificar encabezados del documento: {str(e)}")
        return docx_path, None  # Devolver el documento original si hay error

async def add_page_headers_to_pdf(pdf_path, base_code):
    """
    Modifica un PDF para añadir encabezados diferentes a cada página
    con el formato exacto base_code_Part1, base_code_Part2, etc.
    Mantiene el código base exactamente como está, sin modificarlo.
    """
    try:
        # Crear un nuevo PDF con los encabezados correctos
        output_pdf = os.path.join(os.path.dirname(pdf_path), f"headers_{os.path.basename(pdf_path)}")
        
        # Abrir el PDF original
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        # Para cada página, añadir el encabezado correcto
        for i, page in enumerate(reader.pages):
            # Crear un PDF en memoria con el encabezado
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=letter)
            
            # Dibujar un rectángulo blanco para cubrir completamente cualquier encabezado existente
            can.setFillColorRGB(1, 1, 1)  # Color blanco
            # Bajar el rectángulo y el texto unos 20 puntos
            can.rect(0, 750, 612, 28, fill=True, stroke=False)  # 750 en vez de 770
            
            # Configurar el encabezado con el número de parte correcto
            part_number = i + 1
            header_text = f"{base_code}_Part{part_number}"
            
            # Añadir el texto del encabezado en la posición correcta (esquina superior izquierda, pero más abajo)
            can.setFillColorRGB(0, 0, 0)  # Color negro para el texto
            can.setFont("Helvetica", 10)
            can.drawString(25, 765, header_text)  # 765 en vez de 785
            can.save()
            
            # Mover al inicio del BytesIO
            packet.seek(0)
            watermark = PdfReader(packet)
            
            # Fusionar la página original con el encabezado
            page.merge_page(watermark.pages[0])
            writer.add_page(page)
            
            logger.info(f"Añadido encabezado a página {part_number}: {header_text}")
        
        # Guardar el PDF modificado
        with open(output_pdf, "wb") as output_stream:
            writer.write(output_stream)
        
        logger.info(f"PDF con encabezados modificados guardado en: {output_pdf}")
        
        # Reemplazar el PDF original con el modificado
        shutil.move(output_pdf, pdf_path)
        
        return pdf_path
    
    except Exception as e:
        logger.error(f"Error al añadir encabezados al PDF: {str(e)}")
        return None

async def convert_to_pdf(docx_path, output_dir):
    """
    Convierte un documento Word a PDF usando LibreOffice de manera simple.
    """
    try:
        # Nombre base del archivo sin extensión
        base_name = Path(docx_path).stem
        
        # Comando simple para convertir a PDF
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ]
        
        logger.info(f"Ejecutando: {' '.join(cmd)}")
        
        # Ejecutar el comando
        process = subprocess.run(cmd, capture_output=True, text=True)
        
        # Registrar la salida
        if process.stdout:
            logger.info(f"Salida: {process.stdout}")
        if process.stderr:
            logger.warning(f"Error: {process.stderr}")
        
        # Verificar el archivo PDF generado
        expected_pdf = os.path.join(output_dir, f"{base_name}.pdf")
        
        if os.path.exists(expected_pdf):
            return expected_pdf
        else:
            # Listar archivos en el directorio para diagnóstico
            files = os.listdir(output_dir)
            logger.info(f"Archivos en directorio: {files}")
            
            # Buscar cualquier PDF generado
            for file in files:
                if file.endswith(".pdf") and file.startswith(Path(docx_path).name.split("_")[0]):
                    pdf_path = os.path.join(output_dir, file)
                    logger.info(f"PDF encontrado: {pdf_path}")
                    return pdf_path
            
            logger.error("No se encontró ningún PDF generado")
            return None
            
    except Exception as e:
        logger.error(f"Error en conversión: {str(e)}")
        return None

@app.get("/", summary="Información de la API")
async def root():
    """
    Retorna información básica sobre la API.
    """
    return {
        "mensaje": "API de conversión de Word a PDF",
        "descripcion": "Esta API permite convertir documentos Word a formato PDF",
        "uso": "Envía un archivo Word mediante POST a /convert/",
        "documentacion": "/docs",
        "version": "1.0.0"
    }

@app.get("/health", summary="Verificación de estado")
async def health_check():
    """
    Endpoint para verificar el estado del servicio.
    """
    return {"status": "ok", "message": "El servicio está funcionando correctamente"}

if __name__ == "__main__":
    import uvicorn
    
    # Determinar el puerto desde la variable de entorno o usar 8080 por defecto
    port = int(os.environ.get("PORT", 8080))
    
    logger.info(f"Iniciando servidor en el puerto {port}")
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)
